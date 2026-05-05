"""
job_agent.py — Job Application Agent
Pauses at each step for user approval before proceeding.

Usage:
    python job_agent.py                          # interactive board picker
    python job_agent.py --board builtin          # go straight to builtin.com
    python job_agent.py --board governmentjobs   # go straight to governmentjobs.com
    python job_agent.py --board indeed           # go straight to indeed.com
    python job_agent.py --url "https://..."      # skip board picker, use direct URL

Requirements:
    pip install playwright python-docx pyyaml ollama
    playwright install chromium
"""

import argparse
import copy
import json
import os
import platform
import re
import shutil
import subprocess
import sys
import tempfile
from datetime import datetime
from pathlib import Path

import ollama
import yaml
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from playwright.sync_api import sync_playwright
from job_agent_support.db import already_applied, log_application

# ── Config ────────────────────────────────────────────────────────────────────
COMPONENTS_PATH  = Path("data/job-apps/components.yaml")
OUTPUT_DIR       = Path("data/job-apps/output")
ACTIONS_LOG      = Path("logs/actions_log.json")
CREDENTIALS_PATH = Path("job_agent_support/credentials.yaml")
DB_PATH          = Path("data/job-apps/applications.db")
CHAT_MODEL       = "qwen3:8b"

OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
ACTIONS_LOG.parent.mkdir(parents=True, exist_ok=True)

# ── Logging ───────────────────────────────────────────────────────────────────

def _strip_llm_raw(text: str) -> str:
    text = re.sub(r'<think>[\s\S]*?</think>', '', text, flags=re.IGNORECASE).strip()
    text = re.sub(r'^```json\s*', '', text, flags=re.MULTILINE)
    text = re.sub(r'^```\s*',     '', text, flags=re.MULTILINE)
    text = re.sub(r'\s*```$',     '', text)
    return text.strip()


def _add_hyperlink(para, url: str, text: str, size_pt: float = 11):
    part = para.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(size_pt * 2)))
    rPr.append(sz)
    run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    run.append(t)
    hyperlink.append(run)
    para._p.append(hyperlink)


def log_action(action: str, job_title: str, company: str,
               outcome: str, notes: str = ""):
    entry = {
        "timestamp": datetime.now().strftime("%Y-%m-%dT%H:%M:%S"),
        "action":    action,
        "job_title": job_title,
        "company":   company,
        "outcome":   outcome,
        "notes":     notes,
    }
    entries = []
    if ACTIONS_LOG.exists():
        with open(ACTIONS_LOG, encoding="utf-8") as f:
            try:
                entries = json.load(f)
            except json.JSONDecodeError:
                entries = []
    entries.append(entry)
    with open(ACTIONS_LOG, "w", encoding="utf-8") as f:
        json.dump(entries, f, indent=2)
    print(f"  📋 Logged: {action} → {outcome}")


# ── User approval ─────────────────────────────────────────────────────────────

def ask_approval(prompt: str, details: str = "") -> bool:
    print()
    print("=" * 60)
    print(f"⏸  APPROVAL REQUIRED: {prompt}")
    if details:
        print()
        print(details)
    print("=" * 60)
    while True:
        answer = input("  Proceed? [y/n/q to quit]: ").strip().lower()
        if answer == "y":
            return True
        elif answer == "n":
            return False
        elif answer == "q":
            print("Exiting.")
            sys.exit(0)
        else:
            print("  Please enter y, n, or q.")


def ask_choice(prompt: str, options: dict, details: str = "") -> str:
    """Present a labeled menu of choices and return the selected key."""
    print()
    print("=" * 60)
    print(f"⏸  {prompt}")
    if details:
        print()
        print(details)
    print()
    for key, label in options.items():
        print(f"  [{key}] {label}")
    print(f"  [q] Quit")
    print("=" * 60)
    valid = set(options.keys())
    while True:
        answer = input("  Choice: ").strip().lower()
        if answer == "q":
            print("Exiting.")
            sys.exit(0)
        if answer in valid:
            return answer
        print(f"  Please enter one of: {', '.join(sorted(valid))}, q")


# ── File opener (cross-platform) ──────────────────────────────────────────────

def _open_file(path):
    try:
        system = platform.system()
        if system == "Windows":
            os.startfile(str(path))
        elif system == "Darwin":
            subprocess.run(["open", str(path)])
        else:
            subprocess.run(["xdg-open", str(path)])
    except Exception as e:
        print(f"  Note: could not auto-open file ({e}) — open it manually.")


# ── Page counter ─────────────────────────────────────────────────────────────

def get_page_count(docx_path: Path) -> int:
    """Return the page count of a .docx file.

    Uses LibreOffice headless if available; otherwise estimates from line count.
    """
    if shutil.which("soffice"):
        try:
            with tempfile.TemporaryDirectory() as tmp:
                subprocess.run(
                    ["soffice", "--headless", "--convert-to", "pdf",
                     "--outdir", tmp, str(docx_path)],
                    capture_output=True, timeout=30, check=True,
                )
                pdf_path = Path(tmp) / (docx_path.stem + ".pdf")
                if pdf_path.exists():
                    data = pdf_path.read_bytes()
                    # /Count N in the PDF Pages tree holds total page count
                    counts = re.findall(rb'/Count\s+(\d+)', data)
                    if counts:
                        return max(int(c) for c in counts)
        except Exception:
            pass

    # Heuristic fallback: estimate lines at 11pt within 0.75" margins (~47 lines/page)
    doc = Document(str(docx_path))
    lines = sum(max(1, (len(p.text) + 89) // 90) for p in doc.paragraphs if p.text.strip())
    return max(1, (lines + 46) // 47)


# ── Credentials loader ────────────────────────────────────────────────────────

def load_credentials() -> dict:
    if not CREDENTIALS_PATH.exists():
        print(f"\n  ⚠️  credentials.yaml not found at {CREDENTIALS_PATH}")
        print("  Creating a blank one — fill it in before re-running.\n")
        CREDENTIALS_PATH.parent.mkdir(parents=True, exist_ok=True)
        with open(CREDENTIALS_PATH, "w") as f:
            f.write("builtin:\n  email: \"\"\n  password: \"\"\n\n"
                    "governmentjobs:\n  email: \"\"\n  password: \"\"\n\n"
                    "indeed:\n  email: \"\"\n  password: \"\"\n")
        sys.exit(0)
    with open(CREDENTIALS_PATH, encoding="utf-8") as f:
        return yaml.safe_load(f)


# ── Board picker ──────────────────────────────────────────────────────────────

def pick_board() -> str:
    print("\nWhich job board would you like to use?")
    print("  1. builtin.com")
    print("  2. governmentjobs.com")
    print("  3. indeed.com")
    print("  4. Enter a direct job URL")
    while True:
        choice = input("Choice [1-4]: ").strip()
        if choice == "1": return "builtin"
        if choice == "2": return "governmentjobs"
        if choice == "3": return "indeed"
        if choice == "4": return "url"
        print("  Please enter 1, 2, 3, or 4.")


def load_board_module(board: str):
    if board == "builtin":
        from job_agent_support.boards.builtin import login, browse_jobs, get_job_listings, go_to_next_page
    elif board == "governmentjobs":
        from job_agent_support.boards.governmentjobs import login, browse_jobs, get_job_listings, go_to_next_page
    elif board == "indeed":
        from job_agent_support.boards.indeed import login, browse_jobs, get_job_listings, go_to_next_page
    else:
        raise ValueError(f"Unknown board: {board}")
    return {"login": login, "browse_jobs": browse_jobs,
            "get_job_listings": get_job_listings, "go_to_next_page": go_to_next_page}


# ── ATS detection ─────────────────────────────────────────────────────────────

def detect_ats(url: str) -> str:
    if "greenhouse.io" in url or "boards.greenhouse.io" in url:
        return "greenhouse"
    if "lever.co" in url:
        return "lever"
    if "myworkdayjobs" in url or "workday.com" in url:
        return "workday"
    if "icims.com" in url:
        return "icims"
    return "generic"


def load_ats_module(ats: str):
    if ats == "greenhouse":
        from job_agent_support.ats.greenhouse import fill_page, has_next_page, click_next, click_submit
    elif ats == "lever":
        from job_agent_support.ats.lever import fill_page, has_next_page, click_next, click_submit
    elif ats == "workday":
        from job_agent_support.ats.workday import fill_page, has_next_page, click_next, click_submit
    elif ats == "icims":
        from job_agent_support.ats.icims import fill_page, has_next_page, click_next, click_submit
    else:
        return None
    return {"fill_page": fill_page, "has_next_page": has_next_page,
            "click_next": click_next, "click_submit": click_submit}


# ── Easy Apply detection ───────────────────────────────────────────────────────

def detect_easy_apply(page) -> dict | None:
    """Return {url} if an Easy Apply button is visible on the current page, else None."""
    selectors = [
        'a:has-text("Easy Apply")',
        'button:has-text("Easy Apply")',
        '[data-testid="easy-apply"]',
        '.ia-IndeedApplyButton',
        'button:has-text("Easily apply")',
        '[data-testid="indeedApplyButton"]',
    ]
    for sel in selectors:
        try:
            el = page.query_selector(sel)
            if el and el.is_visible():
                href = el.get_attribute("href")
                return {"url": href or page.url}
        except Exception:
            pass
    return None


# ── Components loader ─────────────────────────────────────────────────────────

def load_components() -> dict:
    if not COMPONENTS_PATH.exists():
        print(f"ERROR: components.yaml not found at {COMPONENTS_PATH}")
        sys.exit(1)
    with open(COMPONENTS_PATH, encoding="utf-8") as f:
        return yaml.safe_load(f)


# ── Step 1: Read job description ──────────────────────────────────────────────

def read_job_description(url: str, page) -> dict:
    print(f"\n{'='*60}")
    print("STEP 1 — Reading job description")
    print(f"  URL: {url}")
    print(f"{'='*60}")

    print("  Navigating to job posting...")
    page.goto(url, wait_until="domcontentloaded", timeout=30000)
    page.wait_for_timeout(2000)
    full_text = page.inner_text("body")
    title_tag = page.title()

    print("  Extracting job details with LLM...")
    prompt = f"""Extract the following from this job posting text.
Return ONLY a JSON object with keys: job_title, company, summary, requirements, responsibilities, pay, address.
Keep each value concise (under 300 words each).
- pay: salary range or hourly rate as a plain string; "Not listed" if absent
- address: office city/state, "Remote", or "Hybrid – [city]"; null if absent

Job posting text:
{full_text[:6000]}

Return only valid JSON, no markdown, no explanation."""

    response = ollama.chat(
        model=CHAT_MODEL,
        messages=[{"role": "user", "content": prompt}],
        options={"temperature": 0},
    )
    raw = _strip_llm_raw(response["message"]["content"])

    try:
        job_info = json.loads(raw)
    except json.JSONDecodeError:
        job_info = {
            "job_title":        title_tag,
            "company":          "Unknown",
            "summary":          full_text[:500],
            "requirements":     "",
            "responsibilities": "",
            "pay":              "Not listed",
            "address":          None,
        }

    job_info["url"]       = url
    job_info["full_text"] = full_text[:8000]

    # Find apply URL
    apply_url = page.evaluate("""() => {
        const links = [...document.querySelectorAll('a')];
        const applyLink = links.find(a => /apply/i.test(a.innerText) && a.href && !a.href.includes('builtin'));
        return applyLink ? applyLink.href : null;
    }""")
    job_info["apply_url"] = apply_url or url

    print(f"\n  Job Title : {job_info.get('job_title', 'N/A')}")
    print(f"  Company   : {job_info.get('company', 'N/A')}")
    print(f"  Apply URL : {job_info['apply_url']}")
    print(f"\n  Summary   : {job_info.get('summary', '')[:300]}...")
    return job_info


# ── Step 2: Select components ─────────────────────────────────────────────────

def select_components(job_info: dict, components: dict) -> dict:
    print(f"\n{'='*60}")
    print("STEP 2 — Selecting relevant components")
    print(f"{'='*60}")

    available_skill_keys = list(components.get("skills", {}).keys())
    skill_descriptions = {
        k: v.get("value", "")[:60]
        for k, v in components.get("skills", {}).items()
    }
    all_project_titles = [p["title"] for p in components.get("projects", [])]
    all_experience     = components.get("experience", [])
    cover_para_info    = {
        k: {"position": v.get("position", "body"), "tags": v.get("tags", [])}
        for k, v in components.get("cover_letter_paragraphs", {}).items()
    }

    exp_options = {}
    for exp in all_experience:
        exp_options[exp["company"]] = {
            "title":   exp["title"],
            "dates":   exp["dates"],
            "bullets": [b["text"] for b in exp.get("bullets", [])],
        }

    prompt = f"""You are a resume tailoring assistant.

JOB DESCRIPTION:
Title: {job_info.get('job_title')}
Company: {job_info.get('company')}
Requirements: {job_info.get('requirements', '')}
Responsibilities: {job_info.get('responsibilities', '')}

YOUR TASK:
1. Select which skills to include (choose from EXACTLY these keys, no others):
   {json.dumps(skill_descriptions, indent=2)}

2. For each work experience below, select EXACTLY 3 bullet points to include.
   You MUST include ALL companies — do not skip any.
   If a company has fewer than 3 bullets available, include all of them.
   {json.dumps(exp_options, indent=2)}

3. Select which projects to include (you MUST include all of them):
   {all_project_titles}

4. Select cover letter paragraph keys for one complete cover letter. Rules:
   - All five selected keys MUST share the same cover letter type prefix (e.g. all start with "it_cloud_")
   - Choose the type prefix whose tags best match the job requirements
   - Select the opening, body_1, body_2, body_3, and closing from that same type
   - The result must be exactly 5 keys in this order: opening, body_1, body_2, body_3, closing
   Available paragraphs (key → position, tags):
   {json.dumps(cover_para_info, indent=2)}

5. List 5-10 keywords from the job description to naturally mirror.

6. List anything the job requires that is NOT available in the skills/experience above.

Return ONLY a JSON object with these exact keys:
- selected_skills: list of skill key names (from the available list only)
- selected_experience: list of objects with "company", "title", "dates", "selected_bullets" (list of bullet text strings)
- selected_projects: list of all project titles
- selected_cover_paragraphs: list of paragraph key names
- keywords_to_mirror: list of keywords
- flagged_fields: list of missing requirements

Return only valid JSON, no markdown."""

    response = ollama.chat(
        model=CHAT_MODEL,
        messages=[{"role": "user", "content": prompt}],
        options={"temperature": 0},
    )
    raw = _strip_llm_raw(response["message"]["content"])

    try:
        selected = json.loads(raw)
    except json.JSONDecodeError:
        print("  WARNING: Could not parse LLM selection — using all components as fallback")
        selected = {
            "selected_skills": available_skill_keys,
            "selected_experience": [
                {
                    "company":          e["company"],
                    "title":            e["title"],
                    "dates":            e["dates"],
                    "selected_bullets": [b["text"] for b in e.get("bullets", [])[:3]],
                }
                for e in all_experience
            ],
            "selected_projects":         all_project_titles,
            "selected_cover_paragraphs": ["it_help_desk_opening", "it_help_desk_body_1", "it_help_desk_body_2", "it_help_desk_body_3", "it_help_desk_closing"],
            "keywords_to_mirror":        [],
            "flagged_fields":            [],
        }

    # Discard LLM-hallucinated companies — only keep exact matches from the library
    valid_companies = set(exp_options.keys())
    selected["selected_experience"] = [
        e for e in selected.get("selected_experience", [])
        if e.get("company") in valid_companies
    ]

    # Safety net: ensure ALL experience entries are always included
    selected_companies = {e["company"] for e in selected.get("selected_experience", [])}
    for exp in all_experience:
        if exp["company"] not in selected_companies:
            selected.setdefault("selected_experience", []).append({
                "company":          exp["company"],
                "title":            exp["title"],
                "dates":            exp["dates"],
                "selected_bullets": [b["text"] for b in exp.get("bullets", [])[:3]],
            })

    # Replace any LLM-hallucinated bullets with real ones from the library,
    # then pad to 3 bullets if fewer were selected
    exp_bullet_lookup = {e["company"]: e for e in all_experience}
    for sel_exp in selected["selected_experience"]:
        company      = sel_exp["company"]
        real_bullets = [b["text"] for b in exp_bullet_lookup.get(company, {}).get("bullets", [])]
        real_set     = set(real_bullets)
        # Keep only bullets that exist verbatim in the library
        valid_bullets = [b for b in sel_exp.get("selected_bullets", []) if b in real_set]
        # Pad to 3 with remaining real bullets not already chosen
        for b in real_bullets:
            if len(valid_bullets) >= 3:
                break
            if b not in valid_bullets:
                valid_bullets.append(b)
        sel_exp["selected_bullets"] = valid_bullets

    selected_proj_titles = set(selected.get("selected_projects", []))
    for proj in components.get("projects", []):
        if proj["title"] not in selected_proj_titles:
            selected.setdefault("selected_projects", []).append(proj["title"])

    valid_skills = [k for k in selected.get("selected_skills", [])
                    if k in available_skill_keys]
    selected["selected_skills"] = valid_skills if valid_skills else available_skill_keys[:4]

    print(f"\n  Skills selected     : {selected.get('selected_skills')}")
    print(f"  Experience included : {[e['company'] for e in selected.get('selected_experience', [])]}")
    print(f"  Projects included   : {selected.get('selected_projects')}")
    print(f"  Cover paragraphs    : {selected.get('selected_cover_paragraphs')}")
    if selected.get("flagged_fields"):
        print(f"\n  ⚠️  FLAGGED (not in library): {selected['flagged_fields']}")

    return selected


# ── Step 3: Generate documents ────────────────────────────────────────────────

def _add_section_header(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.underline = True
    run.font.size = Pt(12)


def _add_bullet(doc, text: str):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(1)
    p.add_run(text).font.size = Pt(11)


def build_resume_docx(job_info: dict, selected: dict,
                      components: dict, output_path: Path):
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(0.75)
        section.right_margin  = Inches(0.75)

    p_info = components["personal"]

    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_before = Pt(0)
    name_p.paragraph_format.space_after  = Pt(2)
    run = name_p.add_run(p_info["name"])
    run.bold = True
    run.font.size = Pt(16)

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_before = Pt(0)
    contact_p.paragraph_format.space_after  = Pt(1)
    contact_p.add_run(
        f"{p_info['location']} | {p_info['phone']} | {p_info['email']}"
    ).font.size = Pt(10)

    links_p = doc.add_paragraph()
    links_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    links_p.paragraph_format.space_before = Pt(0)
    links_p.paragraph_format.space_after  = Pt(10)
    _add_hyperlink(links_p, p_info["portfolio"], p_info["portfolio"], size_pt=10)
    links_p.add_run(" | ").font.size = Pt(10)
    _add_hyperlink(links_p, p_info["linkedin"],  p_info["linkedin"],  size_pt=10)

    _add_section_header(doc, "Skills & Qualifications:")
    skills    = components.get("skills", {})
    label_map = {
        "coding_languages":            "Coding Languages",
        "software_engineer_utilities": "Software & Utilities",
        "software_utilities":          "Software & Utilities",
        "backend":                     "Backend",
        "networking":                  "Networking",
        "soft_skills":                 "Soft Skills",
        "math_courses":                "Math Courses",
        "program_event_operations":    "Program & Event Operations",
        "administrative_facility":     "Administrative & Facility Management",
        "health_safety":               "Health, Safety & Compliance",
        "technology_data":             "Technology & Data Tracking",
        "interpersonal_leadership":    "Interpersonal & Leadership Skills",
    }
    for key in selected.get("selected_skills", []):
        if key in skills:
            label = label_map.get(key, key.replace("_", " ").title())
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(1)
            bold_run = p.add_run(label)
            bold_run.bold = True
            bold_run.font.size = Pt(11)
            p.add_run(f": {skills[key]['value']}").font.size = Pt(11)

    _add_section_header(doc, "Certifications:")
    for cert in components.get("certifications", []):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        bold_run = p.add_run(cert["name"])
        bold_run.bold = True
        bold_run.font.size = Pt(11)
        if cert.get("expires"):
            detail = f". Issued by: {cert['issuer']}. Expires: {cert['expires']}."
        else:
            detail = f" ({cert.get('status', 'In Progress')}). Expected: {cert.get('expected', '')}."
        p.add_run(detail).font.size = Pt(11)

    _add_section_header(doc, "Projects:")
    proj_lookup = {proj["title"]: proj for proj in components.get("projects", [])}
    for proj_title in selected.get("selected_projects", []):
        if proj_title not in proj_lookup:
            continue
        proj = proj_lookup[proj_title]
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(1)
        bold_run = p.add_run(proj["role"])
        bold_run.bold = True
        bold_run.font.size = Pt(11)
        p.add_run(f". {proj['title']}. {proj['dates']}.").font.size = Pt(11)
        for b in proj.get("bullets", []):
            _add_bullet(doc, b["text"])
        for link_url in (proj.get("links") or {}).values():
            lp = doc.add_paragraph()
            lp.paragraph_format.left_indent  = Inches(0.25)
            lp.paragraph_format.space_before = Pt(0)
            lp.paragraph_format.space_after  = Pt(1)
            _add_hyperlink(lp, link_url, link_url, size_pt=10)

    _add_section_header(doc, "Education:")
    for edu in components.get("education", []):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(edu["institution"])
        run.bold = True
        run.font.size = Pt(11)
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(0)
        r2 = p2.add_run(edu["dates"])
        r2.italic = True
        r2.font.size = Pt(11)
        p3 = doc.add_paragraph()
        p3.paragraph_format.space_before = Pt(0)
        p3.paragraph_format.space_after  = Pt(1)
        r3 = p3.add_run(edu["degree"])
        r3.italic = True
        r3.font.size = Pt(11)
        if edu.get("courses"):
            kcp = doc.add_paragraph()
            kcp.paragraph_format.space_before = Pt(0)
            kcp.paragraph_format.space_after  = Pt(1)
            kcp.add_run("Key Courses:").font.size = Pt(11)
            for course in edu["courses"]:
                if " - http" in course:
                    name, url = course.split(" - ", 1)
                    bp = doc.add_paragraph(style='List Bullet')
                    bp.paragraph_format.space_before = Pt(0)
                    bp.paragraph_format.space_after  = Pt(1)
                    bp.add_run(name + " — ").font.size = Pt(11)
                    _add_hyperlink(bp, url, url, size_pt=11)
                else:
                    _add_bullet(doc, course)

    RELATED_TAGS = {"it", "tech", "software_engineer", "networking"}
    exp_lookup   = {e["company"]: e for e in components.get("experience", [])}

    def _is_related(company: str) -> bool:
        for b in exp_lookup.get(company, {}).get("bullets", []):
            if set(b.get("tags", [])) & RELATED_TAGS:
                return True
        return False

    def _write_exp_block(sel_exp: dict):
        p1 = doc.add_paragraph()
        p1.paragraph_format.space_before = Pt(8)
        p1.paragraph_format.space_after  = Pt(0)
        r1 = p1.add_run(sel_exp["title"])
        r1.bold = True
        r1.font.size = Pt(11)
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(0)
        r2 = p2.add_run(sel_exp["company"])
        r2.bold = True
        r2.font.size = Pt(11)
        p3 = doc.add_paragraph()
        p3.paragraph_format.space_before = Pt(0)
        p3.paragraph_format.space_after  = Pt(1)
        p3.add_run(sel_exp["dates"]).font.size = Pt(11)
        for bullet in sel_exp.get("selected_bullets", []):
            _add_bullet(doc, bullet)

    # Deduplicate by company name AND job title (preserve first occurrence of each).
    # Checking title catches cases where the LLM varies the company name but keeps
    # the real title (e.g. "Buchanan Street Elementary" vs "...School", both "Technology Aide").
    seen_companies: set[str] = set()
    seen_titles:    set[str] = set()
    unique_exp: list[dict] = []
    for e in selected.get("selected_experience", []):
        company = e["company"]
        title   = e.get("title", "")
        if company in seen_companies or (title and title in seen_titles):
            continue
        seen_companies.add(company)
        if title:
            seen_titles.add(title)
        # Final bullet guard: keep only verbatim library bullets, pad to 3 if needed.
        # This ensures nothing hallucinated ever reaches the document.
        real_bullets = [b["text"] for b in exp_lookup.get(company, {}).get("bullets", [])]
        real_set     = set(real_bullets)
        valid = [b for b in e.get("selected_bullets", []) if b in real_set]
        for b in real_bullets:
            if len(valid) >= 3:
                break
            if b not in valid:
                valid.append(b)
        e["selected_bullets"] = valid
        unique_exp.append(e)

    related_exp    = [e for e in unique_exp if _is_related(e["company"])]
    additional_exp = [e for e in unique_exp if not _is_related(e["company"])]

    if related_exp:
        _add_section_header(doc, "Related Experience")
        for sel_exp in related_exp:
            _write_exp_block(sel_exp)

    if additional_exp:
        _add_section_header(doc, "Additional Experience:")
        for sel_exp in additional_exp[:2]:
            _write_exp_block(sel_exp)

    doc.save(str(output_path))


def get_cover_letter_paragraphs(selected: dict, components: dict,
                                job_info: dict) -> list[str]:
    paras = components.get("cover_letter_paragraphs", {})
    keys  = [k for k in selected.get("selected_cover_paragraphs", []) if k in paras]

    position_order = {"opening": 0, "body_1": 1, "body_2": 2, "body_3": 3, "closing": 4}
    keys.sort(key=lambda k: position_order.get(paras[k].get("position", "body_1"), 1))

    replacements = {
        "[date]":         datetime.now().strftime("%B %d, %Y"),
        "[job_title]":    job_info.get("job_title", ""),
        "[company_name]": job_info.get("company", ""),
        "[location]":     job_info.get("address", "") or "",
    }

    result = []
    for k in keys:
        text = paras[k]["text"].strip()
        for placeholder, value in replacements.items():
            text = text.replace(placeholder, value)
        # Collapse newlines and extra whitespace from YAML block scalars into single spaces
        text = " ".join(text.split())
        result.append(text)
    return result


def build_cover_letter_docx(job_info: dict, selected: dict,
                            components: dict, output_path: Path,
                            paragraphs: list[str]):
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    p_info = components["personal"]

    def _p(text: str, space_after: float = 12):
        para = doc.add_paragraph(text)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(space_after)
        if para.runs:
            para.runs[0].font.size = Pt(11)
        return para

    _p(datetime.now().strftime("%B %d, %Y"))
    _p("Dear Hiring Manager:")

    for para in paragraphs:
        _p(para)

    _p("Sincerely,", space_after=4)
    _p(p_info["name"], space_after=0)

    doc.save(str(output_path))


def generate_documents(job_info: dict, selected: dict, components: dict) -> dict:
    print(f"\n{'='*60}")
    print("STEP 3 — Generating tailored documents")
    print(f"{'='*60}")

    company_clean = re.sub(r'[^\w\s-]', '', job_info.get("company", "Company"))
    company_clean = company_clean.strip().replace(" ", "_")
    date_str      = datetime.now().strftime("%Y-%m-%d")

    resume_path = OUTPUT_DIR / f"resume_{company_clean}_{date_str}.docx"
    cover_path  = OUTPUT_DIR / f"cover_letter_{company_clean}_{date_str}.docx"

    build_resume_docx(job_info, selected, components, resume_path)

    # Trim experience entries one at a time until the resume fits in 2 pages.
    # Non-related (additional) experience is removed first, then related entries
    # if still over the limit.
    exp_lookup = {e["company"]: e for e in components.get("experience", [])}
    _rel_tags  = {"it", "tech", "software_engineer", "networking"}

    def _exp_is_related(company: str) -> bool:
        for b in exp_lookup.get(company, {}).get("bullets", []):
            if set(b.get("tags", [])) & _rel_tags:
                return True
        return False

    trimmed = copy.deepcopy(selected)
    for _ in range(len(selected.get("selected_experience", [])) + 1):
        pages = get_page_count(resume_path)
        if pages <= 2:
            break
        exp_list    = trimmed["selected_experience"]
        non_related = [i for i, e in enumerate(exp_list) if not _exp_is_related(e["company"])]
        if non_related:
            removed = exp_list.pop(non_related[-1])
        elif exp_list:
            removed = exp_list.pop()
        else:
            print("  ⚠️  Resume exceeds 2 pages but no experience entries remain to remove.")
            break
        print(f"  ✂️  Resume is {pages} pages — removing '{removed['company']}' to fit 2 pages.")
        build_resume_docx(job_info, trimmed, components, resume_path)

    cover_paragraphs = get_cover_letter_paragraphs(selected, components, job_info)
    build_cover_letter_docx(job_info, selected, components, cover_path, cover_paragraphs)

    print(f"\n  Resume saved       : {resume_path}")
    print(f"  Cover letter saved : {cover_path}")

    if selected.get("flagged_fields"):
        print(f"\n  ⚠️  Missing from your library — left blank:")
        for f in selected["flagged_fields"]:
            print(f"      - {f}")

    return {
        "resume_path":       str(resume_path),
        "cover_letter_path": str(cover_path),
    }


# ── Apply-URL resolver (follows intermediate company career pages) ─────────────

def resolve_apply_url(url: str, page) -> str:
    """
    Navigate to url. If the page is an intermediate company career site
    (has a visible 'Apply Now' link but no application form fields), follow
    the link and return the final ATS URL. Otherwise return url unchanged.
    """
    try:
        page.goto(url, wait_until="domcontentloaded", timeout=30000)
        page.wait_for_timeout(2000)
    except Exception:
        return url

    # Check whether the page already has form fields (real application form)
    has_form = page.evaluate("""() => {
        const fields = document.querySelectorAll('input:not([type=hidden]), textarea, select');
        return fields.length > 0;
    }""")
    if has_form:
        return url

    # No form — look for a visible "Apply Now" link to the real ATS
    selectors = [
        'a.button.job-apply',
        'a:has-text("Apply Now")',
        'a:has-text("APPLY NOW")',
        'a:has-text("Apply for this job")',
        'a:has-text("Apply for Job")',
    ]
    for sel in selectors:
        try:
            el = page.query_selector(sel)
            if el and el.is_visible():
                href = el.get_attribute("href")
                if href and href != url:
                    print(f"  ↪  Intermediate page detected — following Apply Now:")
                    print(f"     {href}")
                    return href
        except Exception:
            pass
    return url


# ── Step 4: Inspect form ──────────────────────────────────────────────────────

def inspect_application_form(url: str, page) -> list[dict]:
    print(f"\n{'='*60}")
    print("STEP 4 — Inspecting application form")
    print(f"{'='*60}")

    if page.url != url:
        page.goto(url, wait_until="domcontentloaded", timeout=30000)
        page.wait_for_timeout(2000)

    fields = page.evaluate("""() => {
        const results = [];
        const inputs = document.querySelectorAll('input, textarea, select');
        inputs.forEach(el => {
            if (el.type === 'hidden' || el.type === 'submit') return;
            let label = '';
            if (el.id) {
                const lbl = document.querySelector('label[for="' + el.id + '"]');
                if (lbl) label = lbl.innerText.trim();
            }
            if (!label) label = el.placeholder || el.name || el.type || 'unknown';
            results.push({
                label:      label,
                field_type: el.tagName.toLowerCase() + (el.type ? '[' + el.type + ']' : ''),
                required:   el.required,
                name:       el.name || '',
                id:         el.id || '',
            });
        });
        return results;
    }""")

    print(f"\n  Found {len(fields)} form fields:")
    for f in fields[:20]:
        req = "REQUIRED" if f["required"] else "optional"
        print(f"    [{req}] {f['label']} ({f['field_type']})")

    return fields


# ── Step 5: Fill form (page-by-page with approvals) ───────────────────────────

def fill_application_form(url: str, job_info: dict, docs: dict,
                          components: dict, fields: list, page) -> dict:
    print(f"\n{'='*60}")
    print("STEP 5 — Filling application form")
    print("  Browser is visible — you can click in it at any time.")
    print(f"{'='*60}")

    ats      = detect_ats(url)
    ats_mod  = load_ats_module(ats)
    print(f"  ATS detected: {ats}")

    page.goto(url, wait_until="domcontentloaded", timeout=30000)
    page.wait_for_timeout(2000)

    all_filled  = []
    all_flagged = []

    if ats_mod:
        # ── Greenhouse / known ATS: page-by-page loop ─────────────────────────
        page_num = 1
        while True:
            print(f"\n  --- Form page {page_num} ---")
            result = ats_mod["fill_page"](page, job_info, docs, components)
            all_filled.extend(result.get("filled_fields", []))
            all_flagged.extend(result.get("flagged", []))

            if ats_mod["has_next_page"](page):
                if not ask_approval(f"Page {page_num} filled. Ready to click Next?"):
                    print("  Stopped — browser left open. Complete manually.")
                    input("  Press ENTER to close browser when done.")
                    break
                ats_mod["click_next"](page)
                page.wait_for_timeout(2000)
                page_num += 1
            else:
                # Final page
                print(f"\n  ✅ All pages filled ({page_num} page(s) total).")
                break
    else:
        # ── Generic fallback: fill what we can, then pause ────────────────────
        _generic_fill(page, job_info, docs, components, fields, all_filled, all_flagged)
        print("\n  ⏸  Form filled. Review the browser, then return here.")
        input("  Press ENTER when ready to proceed to submit step...")

    return {"filled_fields": all_filled, "flagged": all_flagged}


def _generic_fill(page, job_info, docs, components, fields,
                  filled_summary, flagged):
    """Basic fill for non-Greenhouse forms."""
    p_info     = components["personal"]
    field_data = {
        "first_name": p_info.get("name", "").split()[0],
        "last_name":  " ".join(p_info.get("name", "").split()[1:]),
        "email":      p_info.get("email", ""),
        "phone":      p_info.get("phone", ""),
        "website":    p_info.get("portfolio", ""),
        "linkedin":   p_info.get("linkedin", ""),
        "location":   p_info.get("location", ""),
    }

    for field in fields:
        field_id   = field.get("id", "")
        field_name = field.get("name", "")
        value      = field_data.get(field_id) or field_data.get(field_name, "")
        if not value:
            continue
        try:
            selector = f"#{field_id}" if field_id else f"[name='{field_name}']"
            el = page.query_selector(selector)
            if el and el.is_visible():
                tag = el.evaluate("el => el.tagName.toLowerCase()")
                if tag in ("input", "textarea"):
                    el.fill(value)
                    filled_summary.append({"field": field_id or field_name, "value": value[:50]})
        except Exception as e:
            flagged.append({"id": field_id, "name": field_name, "error": str(e)})

    # Upload files
    for file_key, file_path in [
        ("resume",       docs.get("resume_path")),
        ("cover_letter", docs.get("cover_letter_path")),
    ]:
        if not file_path or not os.path.exists(file_path):
            continue
        try:
            el = page.query_selector(f"input[type='file'][name*='{file_key}'], input[type='file'][id*='{file_key}']")
            if el:
                el.set_input_files(file_path)
                filled_summary.append({"field": file_key, "value": Path(file_path).name})
        except Exception as e:
            flagged.append({"field": file_key, "error": str(e)})


# ── Step 6: Present summary and submit ───────────────────────────────────────

def present_summary(job_info: dict, docs: dict, fill_result: dict, selected: dict):
    print(f"\n{'='*60}")
    print("STEP 6 — APPLICATION SUMMARY FOR REVIEW")
    print(f"{'='*60}")
    print(f"  Job Title   : {job_info.get('job_title')}")
    print(f"  Company     : {job_info.get('company')}")
    print(f"  URL         : {job_info.get('url')}")
    print(f"\n  Resume      : {docs['resume_path']}")
    print(f"  Cover Letter: {docs['cover_letter_path']}")
    print(f"\n  Fields filled  : {len(fill_result.get('filled_fields', []))}")

    flagged = [f for f in fill_result.get("flagged", []) if f.get("name") or f.get("id")]
    if flagged:
        print(f"\n  ⚠️  Items requiring your review:")
        for f in flagged:
            print(f"      - {f.get('name') or f.get('id', str(f))}")


def submit_application(page, job_info: dict):
    """Attempt to click the Submit button."""
    ats     = detect_ats(job_info.get("apply_url", ""))
    ats_mod = load_ats_module(ats)
    if ats_mod:
        ats_mod["click_submit"](page)
    else:
        # Generic submit attempt
        for sel in ['button[type="submit"]', 'input[type="submit"]', 'button:has-text("Submit")']:
            try:
                el = page.query_selector(sel)
                if el and el.is_visible():
                    el.click()
                    page.wait_for_timeout(3000)
                    print("  ✅ Submit clicked.")
                    return
            except Exception:
                pass
        print("  ⚠️  Could not find Submit button — please click it manually.")
        input("  Press ENTER after submitting manually...")


# ── DB logging helper ─────────────────────────────────────────────────────────

def _log_application(job_info: dict, docs: dict, url: str, source_board: str) -> None:
    log_application(DB_PATH, {
        "job_title":    job_info.get("job_title"),
        "company":      job_info.get("company"),
        "job_board":    source_board,
        "date_applied": datetime.now().strftime("%Y-%m-%d"),
        "pay":          job_info.get("pay", "Not listed"),
        "address":      job_info.get("address"),
        "apply_url":    job_info.get("apply_url", url),
        "resume_path":  docs.get("resume_path"),
        "easy_apply":   0,
    })


# ── Per-job pipeline ──────────────────────────────────────────────────────────

def run_application_pipeline(url: str, page, components: dict, source_board: str = "direct"):
    """Run the full Steps 1-6 pipeline for a single job URL."""

    # ── STEP 1 ────────────────────────────────────────────────────────────────
    job_info = read_job_description(url, page)
    job_info["source_board"] = source_board
    log_action("read_job_description",
               job_info.get("job_title", ""), job_info.get("company", ""), "success")

    # ── Easy Apply check ──────────────────────────────────────────────────────
    easy_apply_info = detect_easy_apply(page)
    if easy_apply_info:
        job_info["is_easy_apply"] = True
        easy_url = easy_apply_info["url"]

        print("\n" + "=" * 60)
        print("⚡ EASY APPLY DETECTED")
        print(f"   Job Title : {job_info.get('job_title')}")
        print(f"   Company   : {job_info.get('company')}")
        print(f"   Board     : {source_board}")
        print()
        print("   Easy Apply Link:")
        print(f"   👉  {easy_url}")
        print()
        print("   Generating tailored documents for your reference...")
        print("=" * 60)

        selected = select_components(job_info, components)
        docs     = generate_documents(job_info, selected, components)

        print(f"\n  📄 Opening documents for reference...")
        _open_file(docs["resume_path"])
        _open_file(docs["cover_letter_path"])

        answer = input("\n  Did you apply to this job? [y/n]: ").strip().lower()
        if answer == "y":
            log_application(DB_PATH, {
                "job_title":    job_info.get("job_title"),
                "company":      job_info.get("company"),
                "job_board":    source_board,
                "date_applied": datetime.now().strftime("%Y-%m-%d"),
                "pay":          job_info.get("pay", "Not listed"),
                "address":      job_info.get("address"),
                "apply_url":    easy_url,
                "resume_path":  docs.get("resume_path"),
                "easy_apply":   1,
            })
            log_action("easy_apply_confirmed",
                       job_info.get("job_title", ""), job_info.get("company", ""), "applied")
            print("  ✅ Application logged.")
        return

    # ── Full pipeline path ────────────────────────────────────────────────────

    # ── STEP 2 ────────────────────────────────────────────────────────────────
    selected = select_components(job_info, components)
    log_action("selected_components",
               job_info.get("job_title", ""), job_info.get("company", ""), "success")

    # ── STEP 3 ────────────────────────────────────────────────────────────────
    docs = generate_documents(job_info, selected, components)
    log_action("generated_documents",
               job_info.get("job_title", ""), job_info.get("company", ""), "success",
               f"Resume: {docs['resume_path']}")

    print(f"\n  📄 Opening documents for review...")
    _open_file(docs["resume_path"])
    _open_file(docs["cover_letter_path"])

    if not ask_approval("Review and edit the documents, then confirm to proceed to form inspection."):
        log_action("step_3_approval", job_info.get("job_title", ""), job_info.get("company", ""), "rejected")
        print("  Stopped after Step 3. Documents saved.")
        return

    # ── STEPS 4–6: Inspect → Fill → (Next → Re-inspect →...) → Submit ──────────
    apply_url = resolve_apply_url(job_info.get("apply_url", url), page)
    job_info["apply_url"] = apply_url

    all_filled  = []
    all_flagged = []
    page_num    = 0

    while True:
        page_num += 1

        # ── Step 4: Inspect current page ─────────────────────────────────────
        fields = inspect_application_form(apply_url, page)
        log_action("inspected_form",
                   job_info.get("job_title", ""), job_info.get("company", ""),
                   "success", f"page {page_num}, {len(fields)} fields")

        step4_choice = ask_choice(
            f"Page {page_num} — {len(fields)} field(s) found",
            {
                "y": "Fill this page",
                "r": "Navigate manually — re-inspect when ready",
                "n": "Stop here",
            },
        )
        if step4_choice == "n":
            log_action("step_4_approval",
                       job_info.get("job_title", ""), job_info.get("company", ""), "rejected")
            print("  Stopped. Documents saved.")
            return
        if step4_choice == "r":
            input("\n  Navigate the browser to the application page, then press ENTER to re-inspect.")
            apply_url = page.url
            job_info["apply_url"] = apply_url
            page_num -= 1
            continue

        # ── Step 5: Detect ATS and fill this page ────────────────────────────
        current_url = page.url
        ats     = detect_ats(current_url)
        ats_mod = load_ats_module(ats)
        print(f"\n  ATS detected : {ats}")
        print("  Filling page — browser is visible, you can click in it at any time.")

        if ats_mod:
            result = ats_mod["fill_page"](page, job_info, docs, components)
        else:
            filled_list, flagged_list = [], []
            _generic_fill(page, job_info, docs, components, fields, filled_list, flagged_list)
            result = {"filled_fields": filled_list, "flagged": flagged_list}
            print("\n  Form filled as much as possible — review the browser.")

        all_filled.extend(result.get("filled_fields", []))
        all_flagged.extend(result.get("flagged", []))
        log_action("filled_form",
                   job_info.get("job_title", ""), job_info.get("company", ""),
                   "success", f"page {page_num}, {len(result.get('filled_fields', []))} fields filled")

        # ── Check for next page ───────────────────────────────────────────────
        has_next = ats_mod is not None and ats_mod["has_next_page"](page)

        if has_next:
            next_choice = ask_choice(
                f"Page {page_num} filled — next page detected",
                {
                    "y": "Click Next and re-inspect",
                    "m": "I submitted manually — log and finish",
                    "n": "Stop here without submitting",
                },
            )
            if next_choice == "n":
                return
            if next_choice == "m":
                log_action("final_submission",
                           job_info.get("job_title", ""), job_info.get("company", ""),
                           "manual", "User submitted manually")
                _log_application(job_info, docs, url, source_board)
                print("  ✅ Application logged as manually submitted.")
                return
            ats_mod["click_next"](page)
            page.wait_for_timeout(2000)
            apply_url = page.url
            continue  # back to Step 4 for the next page

        # ── Step 6: Final page — present summary and submit ───────────────────
        present_summary(job_info, docs,
                        {"filled_fields": all_filled, "flagged": all_flagged}, selected)

        submit_choice = ask_choice(
            "Final page reached — how would you like to proceed?",
            {
                "y": "Submit the application now",
                "m": "I already submitted manually — log and finish",
                "n": "Don't submit — keep documents only",
            },
            "Review any flagged items above before submitting.",
        )

        if submit_choice == "n":
            log_action("final_submission",
                       job_info.get("job_title", ""), job_info.get("company", ""),
                       "rejected", "User chose not to submit")
            print("\n  Application NOT submitted. Documents saved.")
            return

        if submit_choice == "m":
            log_action("final_submission",
                       job_info.get("job_title", ""), job_info.get("company", ""),
                       "manual", "User submitted manually")
            _log_application(job_info, docs, url, source_board)
            print("  ✅ Application logged as manually submitted.")
            return

        submit_application(page, job_info)
        log_action("submitted",
                   job_info.get("job_title", ""), job_info.get("company", ""), "success")
        _log_application(job_info, docs, url, source_board)
        print(f"\n✅ Application submitted.")
        print(f"   Resume      : {docs['resume_path']}")
        print(f"   Cover Letter: {docs['cover_letter_path']}")
        print(f"   Log         : {ACTIONS_LOG}")
        break


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Job Application Agent")
    parser.add_argument("--url",   help="Direct job posting URL (skips board picker)")
    parser.add_argument("--board", choices=["builtin", "governmentjobs", "indeed"],
                        help="Job board to browse")
    args = parser.parse_args()

    print("\n" + "="*60)
    print("JOB APPLICATION AGENT")
    print("Pauses at each step for your approval.")
    print("="*60)

    components = load_components()

    # ── Direct URL mode (no board login) ──────────────────────────────────────
    if args.url:
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=False)
            page    = browser.new_page()
            run_application_pipeline(args.url, page, components)
            browser.close()
        return

    # ── Board browse mode ─────────────────────────────────────────────────────
    board = args.board or pick_board()

    if board == "url":
        url = input("  Enter the job URL: ").strip()
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=False)
            page    = browser.new_page()
            run_application_pipeline(url, page, components)
            browser.close()
        return

    credentials = load_credentials()
    board_creds = credentials.get(board, {})

    # Only governmentjobs requires stored credentials (builtin/indeed skip sign-in)
    if board == "governmentjobs":
        if not board_creds.get("email") or not board_creds.get("password"):
            print(f"\n  ⚠️  No credentials found for 'governmentjobs' in {CREDENTIALS_PATH}")
            print("  Fill in your email and password and re-run.")
            sys.exit(1)

    board_mod = load_board_module(board)

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=False)
        page    = browser.new_page()

        # Login
        board_mod["login"](page, board_creds)
        board_mod["browse_jobs"](page)

        # Browse loop
        while True:
            listings_url = page.url  # remember so we can return after each job
            listings     = board_mod["get_job_listings"](page)

            if not listings:
                print("  No listings found on this page.")
            else:
                for listing in listings:
                    print(f"\n{'─'*60}")
                    print(f"  📌  {listing['title']}")
                    print(f"  🏢  {listing['company']}")

                    if already_applied(DB_PATH, url=listing["url"],
                                       job_title=listing["title"],
                                       company=listing["company"]):
                        print("  ⏭  Already applied — skipping.")
                        continue

                    # Navigate to the job detail page for a real preview
                    try:
                        page.goto(listing["url"],
                                  wait_until="domcontentloaded", timeout=30000)
                        page.wait_for_timeout(2000)
                        preview = page.evaluate("""() => {
                            const sel = [
                                '.job-post-item',
                                '[class*="jobDescription"]',
                                '[class*="job-description"]',
                                '[data-testid="jobDescriptionText"]',
                                '[class*="description-content"]',
                            ].join(', ');
                            const el = document.querySelector(sel);
                            const raw = el ? el.innerText : document.body.innerText;
                            return raw.trim().replace(/\\s+/g, ' ').slice(0, 500);
                        }""")
                        print(f"\n{preview}\n")
                    except Exception as exc:
                        print(f"  (Could not load detail page: {exc})")
                        preview = listing.get("snippet", "")
                        if preview:
                            print(f"  {preview[:300]}\n")

                    if not ask_approval(
                        f"Apply to '{listing['title']}' at {listing['company']}?"
                    ):
                        # Return to the listings page and move to the next card
                        page.goto(listings_url,
                                  wait_until="domcontentloaded", timeout=30000)
                        page.wait_for_timeout(2000)
                        continue

                    run_application_pipeline(
                        listing["url"], page, components, source_board=board
                    )

                    # Return to listings after the pipeline finishes
                    page.goto(listings_url,
                              wait_until="domcontentloaded", timeout=30000)
                    page.wait_for_timeout(2000)

            if not ask_approval("Move to the next page of results?"):
                break
            if not board_mod["go_to_next_page"](page):
                print("  No more pages.")
                break

        browser.close()


if __name__ == "__main__":
    main()
