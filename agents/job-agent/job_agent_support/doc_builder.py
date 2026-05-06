"""
job_agent_support/doc_builder.py
Tag-based resume and cover letter builder.

job_agent.py selects a position tag (and up to 2 backups).
This module deterministically picks skills, experience, projects,
and cover letter paragraphs based solely on that tag, then builds
the .docx files.
"""

import re
import subprocess
import tempfile
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Allowed position tags (LLM must choose from this list) ───────────────────

ALLOWED_TAGS: list[str] = [
    "software_engineer",
    "app_dev",
    "web_dev",
    "data_science",
    "it_cloud",
    "it_help_desk",
    "it_network",
    "admin",
    "events",
]


def validate_tag(tag: str) -> bool:
    return tag in ALLOWED_TAGS


# ── Tag → exactly 4 skill keys ────────────────────────────────────────────────

_TAG_SKILLS: dict[str, list[str]] = {
    "software_engineer": ["coding_languages", "software_engineer_utilities", "backend", "soft_skills"],
    "app_dev":           ["coding_languages", "software_engineer_utilities", "backend", "soft_skills"],
    "web_dev":           ["coding_languages", "software_engineer_utilities", "backend", "soft_skills"],
    "data_science":      ["coding_languages", "backend", "technology_data", "soft_skills"],
    "it_cloud":          ["coding_languages", "software_engineer_utilities", "networking", "soft_skills"],
    "it_help_desk":      ["software_engineer_utilities", "networking", "technology_data", "soft_skills"],
    "it_network":        ["networking", "software_engineer_utilities", "technology_data", "soft_skills"],
    "admin":             ["program_event_operations", "administrative_facility", "technology_data", "interpersonal_leadership"],
    "events":            ["program_event_operations", "administrative_facility", "health_safety", "interpersonal_leadership"],
}

# Skill key → section display title
_SKILL_LABELS: dict[str, str] = {
    "coding_languages":            "Coding Languages",
    "software_engineer_utilities": "Software & Utilities",
    "backend":                     "Backend & Cloud",
    "networking":                  "Networking",
    "soft_skills":                 "Soft Skills",
    "math_courses":                "Math Courses",
    "program_event_operations":    "Program & Event Operations",
    "administrative_facility":     "Administrative & Facility Management",
    "health_safety":               "Health, Safety & Compliance",
    "technology_data":             "Technology & Data Tracking",
    "interpersonal_leadership":    "Interpersonal & Leadership",
}

# Tags that include the Projects section on the resume
_TECH_TAGS: frozenset[str] = frozenset({
    "software_engineer", "app_dev", "web_dev",
    "data_science", "it_cloud", "it_help_desk", "it_network",
})

# Tag → experience bullet tags used to classify "related" vs "additional"
_TAG_EXP_TAGS: dict[str, frozenset[str]] = {
    "software_engineer": frozenset({"software_engineer", "it", "tech"}),
    "app_dev":           frozenset({"software_engineer", "it", "tech", "app_dev", "mobile"}),
    "web_dev":           frozenset({"software_engineer", "it", "tech", "web_dev", "frontend"}),
    "data_science":      frozenset({"data_science", "it", "tech", "ml", "data"}),
    "it_cloud":          frozenset({"it", "tech", "cloud"}),
    "it_help_desk":      frozenset({"it", "tech", "help_desk", "troubleshooting"}),
    "it_network":        frozenset({"it", "tech", "networking"}),
    "admin":             frozenset({"admin", "clerical", "purchasing"}),
    "events":            frozenset({"events", "recreation", "coordination", "admin"}),
}

_MONTHS: dict[str, int] = {
    "january": 1, "february": 2, "march": 3,    "april": 4,
    "may": 5,     "june": 6,     "july": 7,      "august": 8,
    "september": 9, "october": 10, "november": 11, "december": 12,
}


# ── Selection helpers ─────────────────────────────────────────────────────────

def select_skills(tag: str, skills: dict) -> list[tuple[str, str]]:
    """Return [(label, value), ...] for exactly 4 skills for this tag."""
    result = []
    for key in _TAG_SKILLS.get(tag, []):
        if key in skills:
            label = _SKILL_LABELS.get(key, key.replace("_", " ").title())
            result.append((label, skills[key]["value"]))
    return result


def _exp_start_date(exp: dict) -> tuple[int, int]:
    dates = exp.get("dates", "")
    start = dates.split("–")[0].split("-")[0].strip().lower()
    parts = start.split()
    try:
        return int(parts[1]), _MONTHS.get(parts[0], 0)
    except (IndexError, ValueError):
        return 0, 0


def _exp_is_related(exp: dict, exp_tags: frozenset) -> bool:
    return any(set(b.get("tags", [])) & exp_tags for b in exp.get("bullets", []))


def select_experience(tag: str, components: dict) -> tuple[list, list]:
    """
    Return (related_exp, additional_exp), both sorted most-recent-first.
    Each entry has a 'selected_bullets' key with up to 3 bullet texts.
    """
    exp_tags = _TAG_EXP_TAGS.get(tag, frozenset())
    related, additional = [], []
    for exp in components.get("experience", []):
        entry = dict(exp)
        entry["selected_bullets"] = [b["text"] for b in exp.get("bullets", [])[:3]]
        (related if _exp_is_related(exp, exp_tags) else additional).append(entry)

    related.sort(key=_exp_start_date, reverse=True)
    additional.sort(key=_exp_start_date, reverse=True)
    return related, additional


def select_cover_paragraphs(tag: str, cover_letter_paragraphs: dict,
                             job_info: dict) -> list[str]:
    """
    Return the 5 assembled cover letter paragraph strings for this tag.
    Falls back to the first backup tag if fewer than 5 paragraphs exist for primary.
    """
    _ORDER = {"opening": 0, "body_1": 1, "body_2": 2, "body_3": 3, "closing": 4}
    replacements = {
        "[date]":         datetime.now().strftime("%B %d, %Y"),
        "[job_title]":    job_info.get("job_title", ""),
        "[company_name]": job_info.get("company", ""),
        "[location]":     job_info.get("address", "") or "",
    }

    keys = sorted(
        [k for k in cover_letter_paragraphs if k.startswith(tag + "_")],
        key=lambda k: _ORDER.get(cover_letter_paragraphs[k].get("position", "body_1"), 1),
    )

    result = []
    for k in keys:
        text = " ".join(cover_letter_paragraphs[k]["text"].strip().split())
        for ph, val in replacements.items():
            text = text.replace(ph, val)
        result.append(text)
    return result


# ── docx utilities ────────────────────────────────────────────────────────────

def _add_hyperlink(para, url: str, text: str, size_pt: float = 11):
    part = para.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    run  = OxmlElement('w:r')
    rPr  = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(size_pt * 2)))
    rPr.append(sz)
    run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    run.append(t)
    hyperlink.append(run)
    para._p.append(hyperlink)


def _section_header(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.underline = True
    run.font.size = Pt(12)


def _bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(1)
    p.add_run(text).font.size = Pt(11)


def _exp_block(doc: Document, exp: dict):
    for i, (text, bold, italic, space_after) in enumerate([
        (exp["title"],   True,  False, 0),
        (exp["company"], True,  False, 0),
        (exp["dates"],   False, False, 1),
    ]):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8) if i == 0 else Pt(0)
        p.paragraph_format.space_after  = Pt(space_after)
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(11)
    for bullet in exp.get("selected_bullets", []):
        _bullet(doc, bullet)


# ── Page counter ──────────────────────────────────────────────────────────────

def get_page_count(docx_path: Path) -> int:
    import shutil
    if shutil.which("soffice"):
        try:
            with tempfile.TemporaryDirectory() as tmp:
                subprocess.run(
                    ["soffice", "--headless", "--convert-to", "pdf",
                     "--outdir", tmp, str(docx_path)],
                    capture_output=True, timeout=30, check=True,
                )
                pdf_path = Path(tmp) / (docx_path.stem + ".pdf")
                if pdf_path.exists():
                    data   = pdf_path.read_bytes()
                    counts = re.findall(rb'/Count\s+(\d+)', data)
                    if counts:
                        return max(int(c) for c in counts)
        except Exception:
            pass
    # Heuristic fallback (~47 lines/page at 11pt in 0.75" margins)
    doc   = Document(str(docx_path))
    lines = sum(max(1, (len(p.text) + 89) // 90) for p in doc.paragraphs if p.text.strip())
    return max(1, (lines + 46) // 47)


# ── Resume builder ────────────────────────────────────────────────────────────

def build_resume_docx(tag: str, job_info: dict, components: dict,
                      output_path: Path, additional_exp: list | None = None):
    """
    Build a tailored resume .docx for the given position tag.

    Pass `additional_exp` to override the additional experience list
    (used by the 2-page fit loop to trim entries one at a time).
    """
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(0.75)
        section.right_margin  = Inches(0.75)

    p_info = components["personal"]

    # ── Header ────────────────────────────────────────────────────────────────
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_before = Pt(0)
    name_p.paragraph_format.space_after  = Pt(2)
    r = name_p.add_run(p_info["name"])
    r.bold = True
    r.font.size = Pt(16)

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_before = Pt(0)
    contact_p.paragraph_format.space_after  = Pt(1)
    contact_p.add_run(
        f"{p_info['location']} | {p_info['phone']} | {p_info['email']}"
    ).font.size = Pt(10)

    links_p = doc.add_paragraph()
    links_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    links_p.paragraph_format.space_before = Pt(0)
    links_p.paragraph_format.space_after  = Pt(10)
    _add_hyperlink(links_p, p_info["portfolio"], p_info["portfolio"], size_pt=10)
    links_p.add_run(" | ").font.size = Pt(10)
    _add_hyperlink(links_p, p_info["linkedin"],  p_info["linkedin"],  size_pt=10)
    if p_info.get("github"):
        links_p.add_run(" | ").font.size = Pt(10)
        _add_hyperlink(links_p, p_info["github"], p_info["github"], size_pt=10)

    # ── Skills (exactly 4, with bold title) ───────────────────────────────────
    _section_header(doc, "Skills & Qualifications:")
    for label, value in select_skills(tag, components.get("skills", {})):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        bold_run = p.add_run(f"{label}: ")
        bold_run.bold = True
        bold_run.font.size = Pt(11)
        p.add_run(value).font.size = Pt(11)

    # ── Certifications (always) ───────────────────────────────────────────────
    _section_header(doc, "Certifications:")
    for cert in components.get("certifications", []):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        r = p.add_run(cert["name"])
        r.bold = True
        r.font.size = Pt(11)
        if cert.get("expires"):
            detail = f". Issued by: {cert['issuer']}. Expires: {cert['expires']}."
        else:
            detail = f" ({cert.get('status', 'In Progress')}). Expected: {cert.get('expected', '')}."
        p.add_run(detail).font.size = Pt(11)

    # ── Projects (tech tags only) ─────────────────────────────────────────────
    if tag in _TECH_TAGS:
        _section_header(doc, "Projects:")
        for proj in components.get("projects", []):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(1)
            r = p.add_run(proj["role"])
            r.bold = True
            r.font.size = Pt(11)
            p.add_run(f". {proj['title']}. {proj['dates']}.").font.size = Pt(11)
            for b in proj.get("bullets", []):
                _bullet(doc, b["text"])
            for link_url in (proj.get("links") or {}).values():
                lp = doc.add_paragraph()
                lp.paragraph_format.left_indent  = Inches(0.25)
                lp.paragraph_format.space_before = Pt(0)
                lp.paragraph_format.space_after  = Pt(1)
                _add_hyperlink(lp, link_url, link_url, size_pt=10)

    # ── Education (always) ────────────────────────────────────────────────────
    _section_header(doc, "Education:")
    for edu in components.get("education", []):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(0)
        r = p.add_run(edu["institution"])
        r.bold = True
        r.font.size = Pt(11)
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(0)
        r2 = p2.add_run(edu["dates"])
        r2.italic = True
        r2.font.size = Pt(11)
        p3 = doc.add_paragraph()
        p3.paragraph_format.space_before = Pt(0)
        p3.paragraph_format.space_after  = Pt(1)
        r3 = p3.add_run(edu["degree"])
        r3.italic = True
        r3.font.size = Pt(11)
        if edu.get("courses"):
            kcp = doc.add_paragraph()
            kcp.paragraph_format.space_before = Pt(0)
            kcp.paragraph_format.space_after  = Pt(1)
            kcp.add_run("Key Courses:").font.size = Pt(11)
            for course in edu["courses"]:
                if " - http" in course:
                    name, url = course.split(" - ", 1)
                    bp = doc.add_paragraph(style="List Bullet")
                    bp.paragraph_format.space_before = Pt(0)
                    bp.paragraph_format.space_after  = Pt(1)
                    bp.add_run(name + " — ").font.size = Pt(11)
                    _add_hyperlink(bp, url, url, size_pt=11)
                else:
                    _bullet(doc, course)

    # ── Experience ────────────────────────────────────────────────────────────
    related_exp, all_additional = select_experience(tag, components)
    add_exp = additional_exp if additional_exp is not None else all_additional

    if related_exp:
        _section_header(doc, "Related Experience:")
        for exp in related_exp:
            _exp_block(doc, exp)

    if add_exp:
        _section_header(doc, "Additional Experience:")
        for exp in add_exp:
            _exp_block(doc, exp)

    doc.save(str(output_path))


# ── Cover letter builder ──────────────────────────────────────────────────────

def _make_cover_letter_doc(p_info: dict, paragraphs: list[str],
                            font_size: float) -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    def _p(text: str, space_after: float = 12):
        para = doc.add_paragraph(text)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(space_after)
        if para.runs:
            para.runs[0].font.size = Pt(font_size)
        return para

    _p(datetime.now().strftime("%B %d, %Y"))
    _p("Dear Hiring Manager:")
    for para in paragraphs:
        _p(para)
    _p("Sincerely,", space_after=4)
    _p(p_info["name"], space_after=0)
    return doc


def build_cover_letter_docx(tag: str, job_info: dict, components: dict,
                             output_path: Path,
                             backup_tags: list[str] | None = None):
    """
    Build and save a cover letter for the given position tag.
    Auto-shrinks font (12pt → 8pt minimum) until the letter fits on 1 page.
    If the primary tag produces < 5 paragraphs, tries backup_tags in order.
    """
    cl_paras = components.get("cover_letter_paragraphs", {})

    # Try primary tag, then backups if cover letter is incomplete
    chosen_tag  = tag
    paragraphs  = select_cover_paragraphs(tag, cl_paras, job_info)
    for bt in (backup_tags or []):
        if len(paragraphs) >= 5:
            break
        print(f"  ℹ️  Tag '{chosen_tag}' has only {len(paragraphs)} cover letter para(s) — trying backup '{bt}'.")
        chosen_tag = bt
        paragraphs = select_cover_paragraphs(bt, cl_paras, job_info)

    p_info    = components["personal"]
    font_size = 12.0

    while True:
        doc = _make_cover_letter_doc(p_info, paragraphs, font_size)
        doc.save(str(output_path))
        if get_page_count(output_path) <= 1 or font_size <= 8:
            break
        font_size -= 1
        print(f"  ✂️  Cover letter > 1 page at {font_size + 1:.0f}pt — retrying at {font_size:.0f}pt.")


# ── Main entry point ──────────────────────────────────────────────────────────

def generate_documents(primary_tag: str, job_info: dict, components: dict,
                       output_dir: Path,
                       backup_tags: list[str] | None = None) -> dict:
    """
    Build resume + cover letter for the given position tag.

    Resume: trims additional experience one entry at a time (oldest first)
    until it fits in 2 pages.
    Cover letter: auto-shrinks font until it fits on 1 page; uses backup
    tags if the primary is missing paragraphs.

    Returns {"resume_path": str, "cover_letter_path": str}.
    """
    company_clean = re.sub(r'[^\w\s-]', '', job_info.get("company", "Company"))
    company_clean = company_clean.strip().replace(" ", "_")
    date_str      = datetime.now().strftime("%Y-%m-%d")

    resume_path = output_dir / f"resume_{company_clean}_{date_str}.docx"
    cover_path  = output_dir / f"cover_letter_{company_clean}_{date_str}.docx"

    _, all_additional = select_experience(primary_tag, components)
    add_exp = list(all_additional)

    # Build resume, trim additional experience until it fits in 2 pages
    build_resume_docx(primary_tag, job_info, components, resume_path,
                      additional_exp=add_exp)
    for _ in range(len(all_additional) + 1):
        if get_page_count(resume_path) <= 2:
            break
        if not add_exp:
            print("  ⚠️  Resume exceeds 2 pages but no additional experience remains to trim.")
            break
        removed = add_exp.pop()
        print(f"  ✂️  Resume > 2 pages — removing '{removed['company']}' from additional experience.")
        build_resume_docx(primary_tag, job_info, components, resume_path,
                          additional_exp=add_exp)

    build_cover_letter_docx(primary_tag, job_info, components, cover_path,
                             backup_tags=backup_tags)

    print(f"\n  Resume saved       : {resume_path}")
    print(f"  Cover letter saved : {cover_path}")

    return {
        "resume_path":       str(resume_path),
        "cover_letter_path": str(cover_path),
    }
